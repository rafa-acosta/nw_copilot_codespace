"""
Microsoft Word document loader.

Handles loading and preprocessing of DOCX files.
"""

from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None

from .base import BaseDataLoader, IngestedData, IngestSource
from .cleaning import TextCleaner
from .utils import validate_file_exists, validate_file_extension


class DocxLoader(BaseDataLoader):
    """
    Loader for Microsoft Word (.docx) documents.
    
    Extracts text from paragraphs, tables, and other document elements.
    Uses python-docx for document parsing.
    
    Example:
        >>> loader = DocxLoader()
        >>> data = loader.load('document.docx')
        >>> print(data.content[:100])
    """
    
    SUPPORTED_EXTENSIONS = ['.docx']
    
    def __init__(self, cleaner: Optional[TextCleaner] = None):
        """
        Initialize DocxLoader.
        
        Args:
            cleaner: Optional TextCleaner instance.
            
        Raises:
            ImportError: If python-docx is not installed.
        """
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install with: pip install python-docx"
            )
        
        self.cleaner = cleaner or TextCleaner()
    
    def validate_source(self, source: str) -> bool:
        """
        Validate that source is a valid DOCX file.
        
        Args:
            source: File path to validate.
            
        Returns:
            True if valid DOCX file, False otherwise.
        """
        if not validate_file_exists(source):
            return False
        return validate_file_extension(source, self.SUPPORTED_EXTENSIONS)
    
    def load(self, source: str, **kwargs) -> IngestedData:
        """
        Load text from a DOCX file.
        
        Args:
            source: Path to the DOCX file.
            **kwargs: Optional parameters:
                - extract_tables (bool): Include table content (default: True).
                
        Returns:
            IngestedData with extracted and cleaned text.
            
        Raises:
            FileNotFoundError: If file does not exist.
            ValueError: If file is not a valid DOCX.
        """
        if not self.validate_source(source):
            raise ValueError(f"Invalid DOCX file: {source}")
        
        extract_tables = kwargs.get('extract_tables', True)
        metadata = {}
        
        try:
            doc = Document(source)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables if requested
            if extract_tables:
                for table_idx, table in enumerate(doc.tables):
                    table_text = f"\n--- Table {table_idx + 1} ---\n"
                    for row in table.rows:
                        row_text = ' | '.join(cell.text for cell in row.cells)
                        table_text += row_text + '\n'
                    text_parts.append(table_text)
            
            raw_text = '\n'.join(text_parts)
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
        except Exception as e:
            raise ValueError(f"Failed to read DOCX {source}: {str(e)}")
        
        # Clean the extracted text
        cleaned_text = self.cleaner.clean(raw_text)
        
        # Create metadata
        ingest_source = self._create_ingest_source(
            source=source,
            source_type='docx',
            metadata=metadata
        )
        
        return IngestedData(content=cleaned_text, source=ingest_source)
